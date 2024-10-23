import os
import tempfile
import streamlit as st
import streamlit_authenticator as stauth
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from docx import Document as DocxDocument
from io import BytesIO
import yaml
from yaml.loader import SafeLoader

# Page config
st.set_page_config(page_title="Hansard Analyzer", layout="wide")

# Initialize session state at the very beginning
if 'authentication_status' not in st.session_state:
    st.session_state['authentication_status'] = None
if 'name' not in st.session_state:
    st.session_state['name'] = None
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'logout' not in st.session_state:
    st.session_state['logout'] = False

def load_config():
    """Load configuration from secrets"""
    try:
        return {
            'credentials': yaml.safe_load(st.secrets["general"]["credentials"]),
            'cookie_name': st.secrets["general"]["cookie_name"],
            'cookie_key': st.secrets["general"]["cookie_key"],
            'cookie_expiry_days': st.secrets["general"]["cookie_expiry_days"],
            'openai_api_key': st.secrets["general"]["OPENAI_API_KEY"]
        }
    except Exception as e:
        st.error(f"Error loading configuration: {str(e)}")
        st.stop()

@st.cache_data(show_spinner=False)
def process_documents(openai_api_key, model_name, uploaded_files, query):
    """Process documents and generate analysis"""
    try:
        # Initialize progress bars
        loading_progress = st.progress(0)
        processing_progress = st.progress(0)

        # Initialize language models
        embeddings = OpenAIEmbeddings(
            model='text-embedding-3-small',
            openai_api_key=openai_api_key
        )
        llm = ChatOpenAI(
            temperature=0,
            model_name=model_name,
            openai_api_key=openai_api_key
        )

        # Process documents
        docs = []
        total_files = len(uploaded_files)
        
        for i, uploaded_file in enumerate(uploaded_files):
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name

            # Load PDF
            loader = PyPDFLoader(file_path=tmp_file_path)
            docs.extend(loader.load())
            
            # Clean up temporary file
            os.remove(tmp_file_path)
            
            # Update progress
            loading_progress.progress((i + 1) / total_files)

        # Split documents
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300
        )
        splits = text_splitter.split_documents(docs)
        processing_progress.progress(0.4)

        # Create vector store
        vectorstore = FAISS.from_documents(splits, embeddings)
        retriever = vectorstore.as_retriever(
            search_kwargs={"k": 5}
        )
        processing_progress.progress(0.6)

        # Create prompt template
        prompt = ChatPromptTemplate.from_template("""
        You are provided with a context extracted from Canadian parliamentary debates (Hansard) concerning various political issues.
        Answer the question by focusing on the relevant party based on the question. Provide the five to six main points and conclusion.
        
        Context: {context}
        Question: {question}
        
        Answer:""")

        # Create chain
        chain = (
            RunnableParallel(
                {"context": retriever, "question": RunnablePassthrough()}
            )
            | prompt
            | llm
            | StrOutputParser()
        )
        
        processing_progress.progress(0.8)

        # Generate response
        response = chain.invoke(query)
        processing_progress.progress(1.0)

        # Create document
        buffer = BytesIO()
        doc = DocxDocument()
        doc.add_paragraph(f"Question: {query}\n\nAnswer:\n")
        doc.add_paragraph(response)
        doc.save(buffer)
        buffer.seek(0)

        # Clear progress bars
        loading_progress.empty()
        processing_progress.empty()

        return response, buffer

    except Exception as e:
        st.error(f"Error processing documents: {str(e)}")
        return None, None

def main():
    config = load_config()
    
    # Create authenticator
    authenticator = stauth.Authenticate(
        config['credentials'],
        config['cookie_name'],
        config['cookie_key'],
        config['cookie_expiry_days']
    )

    # Check if logout was clicked in the previous run
    if st.session_state['logout']:
        st.session_state['logout'] = False
        st.session_state['authentication_status'] = None
        st.session_state['name'] = None
        st.session_state['username'] = None
        for key in list(st.session_state.keys()):
            if key not in ['authentication_status', 'name', 'username', 'logout']:
                del st.session_state[key]
        st.rerun()

    # Show title
    st.title("Hansard Analyzer")

    # Handle Authentication
    if not st.session_state['authentication_status']:
        # Login form
        authentication_status, name, username = authenticator.login()
        
        st.session_state['authentication_status'] = authentication_status
        st.session_state['name'] = name
        st.session_state['username'] = username
        
        if authentication_status is False:
            st.error("Username/password is incorrect")
            return
        elif authentication_status is None:
            st.warning("Please enter your username and password")
            return

    # Main application
    if st.session_state['authentication_status']:
        # Sidebar
        st.sidebar.title("Hansard Analyzer")
        st.sidebar.write(f"Logged in as: {st.session_state['name']}")
        
        # Simple logout button
        if st.sidebar.button('Logout', key='logout_button'):
            st.session_state['logout'] = True
            st.rerun()
        
        st.sidebar.title("Analysis Settings")
        
        # Model selection
        model_name = st.sidebar.selectbox(
            "Select Model",
            ["gpt-3.5-turbo", "gpt-4"]  # Updated model names to match OpenAI's offerings
        )
        
        # File upload
        uploaded_files = st.sidebar.file_uploader(
            "Upload PDF files",
            type="pdf",
            accept_multiple_files=True
        )
        
        # Main content
        st.title("Hansard Insight Analyzer")
        
        # Query input
        query = st.text_input(
            "Enter your query",
            value="What is the position of the Liberal Party on Carbon Pricing?"
        )

        # Analysis button
        if st.button("Analyze Documents"):
            if uploaded_files and query:
                with st.spinner("Processing documents..."):
                    answer, buffer = process_documents(
                        config['openai_api_key'],
                        model_name,
                        uploaded_files,
                        query
                    )
                    if answer and buffer:
                        st.success("Analysis complete!")
                        st.markdown(f"Question: {query}\n\nAnswer: {answer}")
                        st.session_state['buffer'] = buffer
            else:
                st.warning("Please upload PDF files and enter a query.")

        # Download button
        if 'buffer' in st.session_state and st.session_state['buffer'] is not None:
            st.download_button(
                label="Download Analysis",
                data=st.session_state['buffer'],
                file_name="hansard_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
